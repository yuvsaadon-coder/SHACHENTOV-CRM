"""
Builds the single merged CEO handover document from the two source files.

  hafifa-hila.docx        — Hila's handover  (authoritative on conflicts)
  hafifa-mankal-moran.docx — Moran's handover (3 sessions, 2023)

Output: public/templates/hafifa-mankal-meuchad.docx

The merge keeps both documents whole and in order under one cover: Hila's
content leads each subject, Moran's follows as background. Nothing is
summarised away — the point is a complete file, not an abstract.
"""
import re, html, zipfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = 'public/templates/'
OUT = SRC + 'hafifa-mankal-meuchad.docx'

def paragraphs(path):
    x = zipfile.ZipFile(path).read('word/document.xml').decode('utf8')
    x = x.replace('</w:p>', '\n')
    x = re.sub(r'<w:tab/>', '\t', x)
    txt = html.unescape(re.sub(r'<[^>]+>', '', x))
    return [ln.strip() for ln in txt.split('\n')]

def is_heading(line):
    if not line or len(line) > 70:
        return False
    return line.rstrip().endswith(':') or line.startswith('**') or re.match(r'^[א-ת" ]+$', line) and len(line.split()) <= 5

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

def add(text, size=11, bold=False, color=None, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.right_to_left = True
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.rtl = True
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

NAVY = (0x14, 0x13, 0x48)
TEAL = (0x18, 0x9A, 0x9F)

add('עמותת שכן טוב', 14, True, TEAL, align=WD_ALIGN_PARAGRAPH.CENTER)
add('תיק חפיפה — מנכ״ל/ית', 26, True, NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
add('מסמך משולב', 13, False, NAVY, space_before=4, align=WD_ALIGN_PARAGRAPH.CENTER)
add('')
add('המסמך מאחד את שני תיקי החפיפה הקיימים לתפקיד המנכ״ל/ית:', 11, True)
add('חלק א׳ — חפיפת הילה. במקרה של סתירה בין המסמכים, חלק זה הוא הקובע.')
add('חלק ב׳ — חפיפת מורן (שלוש פגישות חפיפה, 2023), כרקע ופירוט מלא.')
add('')
add('שני החלקים מובאים במלואם, ללא קיצור.', 10, False, (0x6B, 0x72, 0x80))

doc.add_page_break()
add('חלק א׳ — חפיפת הילה', 20, True, NAVY)
add('הקובץ הקובע במקרה של סתירה', 10, False, (0x6B, 0x72, 0x80), space_before=2)
add('')
for line in paragraphs(SRC + 'hafifa-hila.docx'):
    if not line:
        continue
    add(line, 13 if is_heading(line) else 11, is_heading(line),
        NAVY if is_heading(line) else None, 10 if is_heading(line) else 0)

doc.add_page_break()
add('חלק ב׳ — חפיפת מורן', 20, True, NAVY)
add('שלוש פגישות חפיפה, 2023 — רקע ופירוט', 10, False, (0x6B, 0x72, 0x80), space_before=2)
add('')
for line in paragraphs(SRC + 'hafifa-mankal-moran.docx'):
    if not line:
        continue
    add(line, 13 if is_heading(line) else 11, is_heading(line),
        NAVY if is_heading(line) else None, 10 if is_heading(line) else 0)

doc.save(OUT)
print('wrote', OUT)
