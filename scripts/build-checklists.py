"""
Builds two downloadable checklist documents, sourced verbatim from the same
content already seeded into knowledgeItems (scripts/seed-knowledge.mjs):
'סדר פעולות חלוקת חג' and 'סדר פעולות הזמנה חצי-שנתית', both derived from
the Jerusalem coordinator handover file.

Output:
  public/templates/hagim/tzeklist-chalukat-chag.docx
  public/templates/hatzi-shnatit/tzeklist-hazmana-hatzi-shnatit.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = (0x14, 0x13, 0x48)
TEAL = (0x18, 0x9A, 0x9F)

HOLIDAY_ITEMS = [
    ('היקף ההזמנה', [
        'ודא שהתקציב הייעודי לחג מאושר ונפרד מהתקציב השוטף',
        'הכן רשימת סניפי ירושלים הקבועים: בקעה, גילה, נווה יעקב, נחלאות, קטמון, קרית מנחם, רמות',
        'הוסף סניפים "אורחים" בירושלים (שלומי באבד, סבתא רבקא, הרב שמעון, מעלה אדומים) — מקבלים יבשים וביצים (12 יח׳) בלבד',
        'הוסף סניפי חוץ (תל אביב, בית שמש, קדימה) — יבשים בלבד; יש להם תקציב נפרד להשלמות',
    ]),
    ('בחירת ספק', [
        'שלח בקשה להצעת מחיר ליבשים לשני ספקים לפחות — אריזות ירושלים וברוכים (ספק החגים בשנים האחרונות: אריזות ירושלים)',
        'ודא שהצעות המחיר כוללות מע״מ לפני ההשוואה',
        'הכנס את הצעות המחיר לקובץ מעקב מחירים רב-שנתי',
        'בבחירת הספק שקול גם אמינות ויכולת לוגיסטית, לא רק מחיר',
        'משא ומתן: אריזות נותנים הנחה קבועה של 7% — חתור ל-9% (חדוה מעבירה לחיים המנהל)',
    ]),
    ('ביצוע ההזמנה', [
        'שלח לספק הנבחר את הכמויות הסופיות יחד עם הכמויות לכל סניף ופרטי הסניפים',
        'שלח כמויות גם לספקים הנוספים: ביצים (בני ס.א.ל), עופות (ינון ברוכים), ירקות, חלות (ברמן)',
        'שלח לכל סניף בפרטי את הכמויות והספקים שלו (אפשר צילום מסך מהאקסל)',
        'מלא את קובץ המעקב והתחשיב הכולל לחג',
        'בדוק מתי צפויה התוצרת של לקט-הרטמן והכווין את הסניפים לזמני חלוקה בהתאם',
    ]),
    ('פרסום', [
        'כחודש לפני החג: אסוף מהסניפים זמני חלוקה ואריזה, תאריכים ומיקומים',
        'העבר את זמני החלוקה והאריזה לצוות הסושיאל לפרסום',
        'בפסח בלבד: אסוף גם זמני ונקודות איסוף חמץ',
        'בפסח בלבד: הזמן מצות',
        'הזמן גלויות חג לראש השנה/פסח עבור החלוקות הגדולות',
    ]),
    ('סגירה', [
        'ודא מול כל סניף שהתיאום מול הספקים תקין ושקיבלו את הכל כראוי',
        'עדכן את קובץ מעקב התקציב בקבלות הספקים לאחר החג',
    ]),
]

SEMIANNUAL_ITEMS = [
    ('היקף ותכנון', [
        'חשב את ההזמנה עבור 24 חלוקות (12 לסניפים דו-שבועיים) — מביא בחשבון ביטולי חלוקות בחגים ועודפים טבעיים',
        'אפשר לסניפים לעדכן את מספר הסלים, בכפוף לאישור הרכז/ת ולתקציב',
        'קח בחשבון את תרומת טפרברג (מיץ ענבים) בחישוב',
    ]),
    ('בחירת ספק', [
        'שלח בקשה להצעת מחיר ליבשים לשני ספקים לפחות (ספק החצי-שנתית בשנים האחרונות: ברוכים)',
        'ודא שהמחירים כוללים מע״מ והכנס אותם לקובץ השוואת המחירים',
        'הכנס את ההצעה לקובץ מעקב מחירים רב-שנתי כדי לעקוב אחרי עליות מחירים',
    ]),
    ('ביצוע ההזמנה', [
        'שלח לספק הנבחר את הכמויות הסופיות + הכמויות לכל סניף + פרטי הסניפים',
        'שלח לכל סניף בפרטי את הכמויות והספקים שלו (אפשר צילום מסך מהאקסל)',
        'אם רכז מבקש הזמנה מוקדמת — ברר אם מדובר בנוחות או בחוסרים אמיתיים; בחוסרים, הבן את הסיבה ומנע הישנות',
        'לקראת שנה חדשה (נובמבר) שלח לינון את טבלת העופות לשנה הבאה (יש פורמט ייעודי)',
    ]),
    ('סגירה', [
        'ודא עם הסניפים שהתיאום מול הספקים תקין ושקיבלו את הכל כראוי',
    ]),
]


def build(title, subtitle, sections, out_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    def add(text, size=11, bold=False, color=None, space_before=0, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.right_to_left = True
        p.paragraph_format.space_before = Pt(space_before)
        if align is not None:
            p.alignment = align
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.font.rtl = True
        if color:
            r.font.color.rgb = RGBColor(*color)
        return p

    add('עמותת שכן טוב', 12, True, TEAL, align=WD_ALIGN_PARAGRAPH.CENTER)
    add(title, 22, True, NAVY, space_before=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add(subtitle, 11, False, (0x6B, 0x72, 0x80), space_before=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add('')

    n = 1
    for section, items in sections:
        add(section, 13, True, NAVY, space_before=10)
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.right_to_left = True
            r = p.add_run(f'☐  {item}')
            r.font.size = Pt(11)
            r.font.rtl = True
            n += 1

    doc.save(out_path)
    print('wrote', out_path)


build(
    'סדר פעולות חלוקת חג',
    'צ׳קליסט מלא לחלוקת מזון לחג — נגזר מקובץ החפיפה של רכז/ת סניפי ירושלים',
    HOLIDAY_ITEMS,
    'public/templates/hagim/tzeklist-chalukat-chag.docx',
)
build(
    'סדר פעולות הזמנה חצי-שנתית',
    'צ׳קליסט להזמנת המלאי החצי-שנתית — נגזר מקובץ החפיפה של רכז/ת סניפי ירושלים',
    SEMIANNUAL_ITEMS,
    'public/templates/hatzi-shnatit/tzeklist-hazmana-hatzi-shnatit.docx',
)
